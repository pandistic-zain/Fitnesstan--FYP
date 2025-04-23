# =============================================================
# Calorie‚ÄëPrediction Trainer  (v6 ‚Äì fixed feature set & ensemble)
# -------------------------------------------------------------
#  Changes (per user request):
#  ‚Ä¢ Revert auto‚Äëfeature detection ‚Äì use only predefined macro + micro
#    nutrients for training (see OPTIONAL_NUM list).
#  ‚Ä¢ Build two models (HistGradientBoosting + Ridge‚ÄëPoly) and their
#    simple **ensemble** (mean of predictions).
#  ‚Ä¢ Produce Markdown + Word (report.docx) with tables & charts.
#  ‚Ä¢ Guidance: deploy the **blend** in Flask by loading both pickles
#    and averaging outputs.
# -------------------------------------------------------------
#  RUN
#      python train_calorie_model.py -d nutrition.csv -o models
#  DEP
#      pip install pandas numpy scikit-learn matplotlib joblib python-docx
# =============================================================
import argparse, logging, pathlib, re, textwrap, joblib
import numpy as np, pandas as pd, matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from sklearn.pipeline import Pipeline
from sklearn.linear_model import Ridge
from sklearn.metrics import mean_absolute_error, r2_score, mean_squared_error
from sklearn.ensemble import HistGradientBoostingRegressor
from sklearn.inspection import permutation_importance
from docx import Document
from docx.shared import Inches

# ---------- logging ---------
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s %(levelname)s %(message)s",
                    datefmt="%Y-%m-%d %H:%M:%S")
logging.getLogger("matplotlib").setLevel(logging.WARNING)
log = logging.getLogger("calorie_trainer")

# ---------- cli -------------
parser = argparse.ArgumentParser()
parser.add_argument("-d", "--data-path", type=pathlib.Path, default=pathlib.Path.cwd()/"nutrition.csv")
parser.add_argument("-o", "--outdir",   type=pathlib.Path, default=pathlib.Path.cwd()/"models")
args = parser.parse_args(); args.outdir.mkdir(parents=True, exist_ok=True)

# ---------- constants -------
MANDATORY_NUM = ["protein", "carbohydrate", "total_fat", "serving_weight", "calories"]
OPTIONAL_NUM  = ["saturated_fat", "fiber", "sugar", "sodium"]
re_digits = re.compile(r"[^0-9.]")

# ---------- helpers ---------
def to_float(series: pd.Series) -> pd.Series:
    """Strip non‚Äënumeric chars & convert to float; preserves index length."""
    return pd.to_numeric(series.astype(str).str.replace(re_digits, "", regex=True), errors="coerce")

# ---------- load ------------
log.info(f"Loading {args.data_path}")
if not args.data_path.exists():
    raise SystemExit("‚ùå Data file not found ‚Üí " + str(args.data_path))
raw = pd.read_csv(args.data_path)
log.info(f"Raw shape {raw.shape}")

# serving weight from text column
if "serving_size" not in raw.columns:
    raise SystemExit("‚ùå 'serving_size' column missing")
raw["serving_weight"] = to_float(raw["serving_size"])

# ensure all feature columns exist and are parsed
for col in MANDATORY_NUM + OPTIONAL_NUM:
    if col in raw.columns:
        raw[col] = to_float(raw[col])
    else:
        raw[col] = pd.Series(np.nan, index=raw.index)

# category fallback
if "category" not in raw.columns:
    raw["category"] = "unknown"

# ---------- filter / impute -------------------------------
clean = raw[MANDATORY_NUM + OPTIONAL_NUM + ["category"]].copy()
clean = clean.dropna(subset=MANDATORY_NUM)
clean[OPTIONAL_NUM] = clean[OPTIONAL_NUM].fillna(0.0)
log.info(f"Clean shape {clean.shape}")

# ---------- design matrix ---------------------------------
X_df = pd.get_dummies(clean[[*MANDATORY_NUM[:-1], *OPTIONAL_NUM, "category"]],
                      columns=["category"], drop_first=True)
feature_names = X_df.columns.tolist()
X = X_df.values
y = clean["calories"].values

X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=0.2, random_state=42)

# ---------- hyper‚Äëgrids -----------------------------------
hgb_grid = {
    "learning_rate":     [0.01, 0.03, 0.05, 0.1, 0.2, 0.3],
    "max_depth":         [None, 3, 5, 7, 9, 12],
    "min_samples_leaf":  [5, 10, 20, 50, 100],
    "l2_regularization": [0.0, 0.01, 0.1, 1.0, 5.0]
}
alpha_vals = [1e-4, 1e-3, 0.01, 0.1, 1, 5, 10, 50, 100, 500, 1000]

# ---------- HGB -------------------------------------------
log.info("üîç Tuning HGB ‚Ä¶")

# Perform grid search to tune HGB
grid_search = GridSearchCV(HistGradientBoostingRegressor(random_state=42), hgb_grid,
                           cv=5, scoring="neg_root_mean_squared_error", n_jobs=12)
grid_search.fit(X_tr, y_tr)
HGB = grid_search.best_estimator_  # Get the best model after grid search

# Predict using the best estimator
pred_hgb = HGB.predict(X_te)

# Get the cross-validation results for HGB from the GridSearchCV object
cv_results = grid_search.cv_results_  # Accessing cv_results_ from GridSearchCV

# Now we need to plot correctly for learning_rate and handle cross-validation results
mean_test_scores = cv_results["mean_test_score"].reshape(len(hgb_grid["max_depth"]), -1).mean(axis=1)

# Plot the validation curve for HGB's learning_rate
plt.figure(figsize=(6, 4))
plt.plot(hgb_grid["learning_rate"], -mean_test_scores, marker="o")
plt.xscale("log")
plt.xlabel("Learning Rate (Log Scale)")
plt.ylabel("Negative Mean Test Score (RMSE)")
plt.title("HGB Learning Rate vs Performance (Cross-validation)")
plt.grid(True)
plt.tight_layout()
plt.savefig(args.outdir/"cv_hgb_learning_rate.png", dpi=250)
plt.show()

# Example: Plot for max_depth
plt.figure(figsize=(6, 4))
plt.plot(hgb_grid["max_depth"], -cv_results["mean_test_score"], marker="o")
plt.xscale("log")
plt.xlabel("Max Depth")
plt.ylabel("Negative Mean Test Score (RMSE)")
plt.title("HGB Max Depth vs Performance (Cross-validation)")
plt.grid(True)
plt.tight_layout()
plt.savefig(args.outdir/"cv_hgb_max_depth.png", dpi=250)
plt.show()

# Example: Plot for min_samples_leaf
plt.figure(figsize=(6, 4))
plt.plot(hgb_grid["min_samples_leaf"], -cv_results["mean_test_score"], marker="o")
plt.xscale("log")
plt.xlabel("Min Samples Leaf")
plt.ylabel("Negative Mean Test Score (RMSE)")
plt.title("HGB Min Samples Leaf vs Performance (Cross-validation)")
plt.grid(True)
plt.tight_layout()
plt.savefig(args.outdir/"cv_hgb_min_samples_leaf.png", dpi=250)
plt.show()

# Example: Plot for l2_regularization
plt.figure(figsize=(6, 4))
plt.plot(hgb_grid["l2_regularization"], -cv_results["mean_test_score"], marker="o")
plt.xscale("log")
plt.xlabel("L2 Regularization")
plt.ylabel("Negative Mean Test Score (RMSE)")
plt.title("HGB L2 Regularization vs Performance (Cross-validation)")
plt.grid(True)
plt.tight_layout()
plt.savefig(args.outdir/"cv_hgb_l2_regularization.png", dpi=250)
plt.show()


# ---------- Ridge + Poly ----------------------------------
# Ridge Œ± curve plotting (cross-validation results)
log.info("üîç Tuning Ridge ‚Ä¶")

# Perform grid search and get the best model
grid_search_ridge = GridSearchCV(Pipeline([
                ("poly", PolynomialFeatures(degree=2, include_bias=False)),
                ("scale", StandardScaler()),
                ("ridge", Ridge())]),
                {"ridge__alpha": alpha_vals}, cv=5,
                scoring="neg_root_mean_squared_error", n_jobs=12)

# Fit the model using grid search
grid_search_ridge.fit(X_tr, y_tr)

# Get the best estimator from the grid search
RIDGE = grid_search_ridge.best_estimator_

# Get the cross-validation results for Ridge
cv_results = grid_search_ridge.cv_results_  # Access cv_results_ from GridSearchCV object

# Reshaping cv_results to match the alpha_vals for plotting
mean_test_scores = cv_results["mean_test_score"].reshape(len(alpha_vals), -1).mean(axis=1)

# Plot the validation curve for Ridge's alpha values
plt.figure(figsize=(6, 4))
plt.plot(alpha_vals, -mean_test_scores, marker="o")
plt.xscale("log")
plt.xlabel("Alpha (Log Scale)")
plt.ylabel("Negative Mean Test Score (RMSE)")
plt.title("Ridge Œ± Regularization Curve (Cross-validation)")
plt.grid(True)
plt.tight_layout()
plt.savefig(args.outdir/"cv_ridge.png", dpi=250)
plt.show()

# ---------- Ensemble --------------------------------------
blend_pred = 0.5 * (pred_hgb + pred_ridge)

# ---------- Metrics ---------------------------------------
metric = lambda y,p: (np.sqrt(mean_squared_error(y,p)), mean_absolute_error(y,p), r2_score(y,p))
rmse_h, mae_h, r2_h = metric(y_te, pred_hgb)
rmse_r, mae_r, r2_r = metric(y_te, pred_ridge)
rmse_b, mae_b, r2_b = metric(y_te, blend_pred)
allowed_err = 0.20 * y_te.mean()

# ---------- save models & plots ---------------------------
models_dir = args.outdir
joblib.dump(HGB,   models_dir/"hgb.pkl")
joblib.dump(RIDGE, models_dir/"ridge_poly.pkl")

perm = permutation_importance(HGB, X_te, y_te, n_repeats=10, random_state=42)
plt.figure(figsize=(max(6, len(feature_names)*0.25),3))
plt.bar(feature_names, perm.importances_mean)
plt.xticks(rotation=90, ha="right", fontsize=6)
plt.tight_layout(); plt.ylabel("Mean ŒîRMSE"); plt.title("HGB Feature Importance")
plt.savefig(models_dir/"importance_hgb.png", dpi=250)

plt.figure(figsize=(4,3))
plt.plot(alpha_vals, -RIDGE.named_steps["ridge"].alpha if hasattr(RIDGE.named_steps["ridge"], "alpha") else alpha_vals, marker="o")
plt.xscale("log"); plt.xlabel("alpha"); plt.ylabel("CV curve")
plt.tight_layout(); plt.title("Ridge Œ± curve"); plt.grid(True)
plt.savefig(models_dir/"cv_ridge.png", dpi=250)

# ---------- Comparison Graph: Model vs FDA Tolerance -----------------
plt.figure(figsize=(6, 4))
plt.bar(["MAE", "RMSE", "FDA 20%"], [mae_b, rmse_b, allowed_err])
plt.ylabel("Calories"); plt.title("Blend vs FDA tolerance"); plt.tight_layout()
plt.savefig(models_dir/"metrics_compare.png", dpi=250)

# ---------- report (markdown) -----------------------------
md = f"""# Calorie Predictor Report\n\n## Features Used\nMandatory: {', '.join(MANDATORY_NUM)}\n\nOptional: {', '.join(OPTIONAL_NUM)}\n\n## Results\n|Model|RMSE|MAE|R¬≤|\n|---|---|---|---|\n|HGB|{rmse_h:.2f}|{mae_h:.2f}|{r2_h:.3f}|\n|Ridge|{rmse_r:.2f}|{mae_r:.2f}|{r2_r:.3f}|\n|Blend|{rmse_b:.2f}|{mae_b:.2f}|{r2_b:.3f}|\n\n## FDA Comparison\n![FDA](metrics_compare.png)\n"""
(models_dir/"report.md").write_text(md, encoding="utf-8")

# ---------- Word doc (brief) ------------------------------
doc = Document()
doc.add_heading('Calorie Predictor Report', 0)
doc.add_heading('1. Features Used', level=1)
doc.add_paragraph(f"Mandatory: {', '.join(MANDATORY_NUM)}")
doc.add_paragraph(f"Optional: {', '.join(OPTIONAL_NUM)}")

doc.add_heading('2. Results', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'RMSE'
hdr_cells[2].text = 'MAE'
hdr_cells[3].text = 'R¬≤'

row = table.add_row().cells
row[0].text = 'HGB'
row[1].text = f"{rmse_h:.2f}"
row[2].text = f"{mae_h:.2f}"
row[3].text = f"{r2_h:.3f}"

row = table.add_row().cells
row[0].text = 'Ridge'
row[1].text = f"{rmse_r:.2f}"
row[2].text = f"{mae_r:.2f}"
row[3].text = f"{r2_r:.3f}"

row = table.add_row().cells
row[0].text = 'Blend'
row[1].text = f"{rmse_b:.2f}"
row[2].text = f"{mae_b:.2f}"
row[3].text = f"{r2_b:.3f}"

doc.add_heading('3. FDA Comparison', level=1)
doc.add_paragraph('The following graph shows the comparison between model metrics and FDA\'s 20% calorie-label tolerance.')
doc.add_picture(str(models_dir/"metrics_compare.png"))

# Add feature importance graph to doc
doc.add_heading('4. Feature Importance', level=1)
doc.add_paragraph("This graph shows the feature importance as determined by the HistGradientBoosting model.")
doc.add_picture(str(models_dir/"importance_hgb.png"))

# Save the Word document
doc.save(models_dir/"report.docx")

log.info(f"‚úÖ Report saved to {models_dir/'report.docx'}")
